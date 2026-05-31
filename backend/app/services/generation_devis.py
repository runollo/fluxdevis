"""Service de generation de devis Word.

Devis oriente CLIENT : presentation commerciale et professionnelle.
Principe (cf. pratique des agences web) :
- la creation est presentee comme UN livrable a UN prix (pas de prix par option,
  source de friction) ; le detail du perimetre est liste sans prix ;
- les avantages (articles offerts + remise) sont mis en valeur pour materialiser
  la valeur percue ;
- l'abonnement mensuel est nettement separe (engagement recurrent distinct) ;
- ancrage de prix : valeur catalogue -> avantages -> net a payer.
"""

from decimal import Decimal, ROUND_HALF_UP
from fractions import Fraction
from io import BytesIO
from pathlib import Path

from app.services.echeances import repartir_au_centime

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.shared import Cm

from app.services.word_helpers import (
    C_NAVY, C_TEXT, C_WHITE, C_AHEAD, C_PAID, C_CURR,
    HEX_NAVY, HEX_PAID,
    fmt_eur, setup_page, force_arial, tbl_no_spacing, full_tbl_borders,
    cell_bg, cell_w, row_height, p_fmt, run, cell_text, hline, spacer,
)

D = Decimal
TVA_RATE = D("0.20")
HEX_AVANTAGE = "EEF7F1"   # vert tres clair (encadre avantages)
_ASSETS = Path(__file__).resolve().parent.parent / "templates" / "assets"

# Socle commun a toutes les offres (elements de reassurance).
SOCLE_COMMUN = [
    ("Hébergement", "Inclus pendant toute la durée de l'abonnement"),
    ("Nom de domaine", "Reprise de l'existant ou création (.com / .fr selon disponibilité) "
                       "+ une adresse e-mail. Renouvellement inclus dans la maintenance"),
    ("Certificat SSL", "Sécurisation HTTPS incluse"),
    ("Design responsive", "Adapté mobile / tablette / desktop"),
    ("Formulaire de contact", "Avec protection anti-spam"),
]


def _q(val) -> Decimal:
    return Decimal(str(val or 0)).quantize(Decimal("0.01"), rounding=ROUND_HALF_UP)


def _fmt_pct(pct) -> str:
    """Formate un pourcentage a la francaise sans zeros inutiles (30 %, 37,5 %)."""
    d = Decimal(str(pct or 0))
    d = d.quantize(Decimal("1")) if d == d.to_integral_value() else d.quantize(Decimal("0.1"))
    return format(d, "f").replace(".", ",") + " %"


# Repartition des echeances selon le plan de paiement (comptant).
_PLAN_PARTS = {
    "100%": [("Solde", Fraction(1, 1))],
    "50/50": [("Acompte 50 %", Fraction(1, 2)), ("Solde 50 %", Fraction(1, 2))],
    "33/33/33": [
        ("Acompte 1/3 (a la commande)", Fraction(1, 3)),
        ("2e versement 1/3 (a mi-parcours)", Fraction(1, 3)),
        ("Solde 1/3 (a la livraison)", Fraction(1, 3)),
    ],
    "50/25/25": [
        ("Acompte 50 % (a la commande)", Fraction(1, 2)),
        ("2e versement 25 %", Fraction(1, 4)),
        ("Solde 25 % (a la livraison)", Fraction(1, 4)),
    ],
    "25/25/25/25": [
        ("Acompte 25 % (a la commande)", Fraction(1, 4)),
        ("2e versement 25 %", Fraction(1, 4)),
        ("3e versement 25 %", Fraction(1, 4)),
        ("Solde 25 % (a la livraison)", Fraction(1, 4)),
    ],
}


def repartition_echeances(plan: str | None, total_ttc) -> list[tuple[str, Decimal]]:
    """Repartit un montant TTC selon le plan de paiement (somme exacte au centime)."""
    parts = _PLAN_PARTS.get(plan or "100%", _PLAN_PARTS["100%"])
    montants = repartir_au_centime(total_ttc, [frac for _, frac in parts])
    return [(label, montants[i]) for i, (label, _) in enumerate(parts)]


# ---------------------------------------------------------------------------
# Calcul des donnees derivees (valeur catalogue, avantages, mensuel net)
# ---------------------------------------------------------------------------

def _calcul_synthese(devis) -> dict:
    """Reconstitue les montants commerciaux a partir du snapshot du devis.

    Renvoie un dict avec : prix_creation_ht, remise_setup, offerts (liste),
    offerts_setup_total, avantages_total, valeur_catalogue,
    mensuel_brut, mensuel_net, remise_pct_recurrent.
    """
    prix_creation_ht = _q(devis.total_ht)
    remise_setup = _q(devis.remise_eur_setup)

    offerts = [
        {"designation": a.designation, "valeur": _q(a.prix_vente)}
        for a in sorted(devis.articles_offerts or [], key=lambda x: x.ordre)
    ]
    offerts_total = sum((o["valeur"] for o in offerts), D("0"))
    offerts_recurrent = _q(getattr(devis, "total_offerts_recurrent_ht", 0) or 0)
    offerts_setup_total = offerts_total - offerts_recurrent
    if offerts_setup_total < 0:
        offerts_setup_total = D("0")

    # Avantages mis en avant cote creation (one-shot) : offerts setup + remise setup.
    avantages_total = offerts_setup_total + remise_setup
    valeur_catalogue = prix_creation_ht + avantages_total

    # Mensuel : brut (catalogue) puis net (apres deduction offerts recurrent et remise).
    mensuel_brut = _q(devis.total_pack_maintenance_ht) + _q(devis.total_options_recurrent_ht)
    remise_pct = _q(devis.remise_pct_recurrent) / D("100")
    mensuel_net = _q((mensuel_brut - offerts_recurrent) * (D("1") - remise_pct))
    if mensuel_net < 0:
        mensuel_net = D("0")

    # Pourcentage d'avantages global (offerts + remise) sur la valeur catalogue.
    avantages_pct = (
        (avantages_total / valeur_catalogue * D("100")).quantize(D("1"), ROUND_HALF_UP)
        if valeur_catalogue > 0 else D("0")
    )

    return {
        "prix_creation_ht": prix_creation_ht,
        "remise_setup": remise_setup,
        "remise_pct_setup": _q(devis.remise_pct_setup),
        "offerts": offerts,
        "offerts_setup_total": offerts_setup_total,
        "avantages_total": avantages_total,
        "avantages_pct": avantages_pct,
        "valeur_catalogue": valeur_catalogue,
        "mensuel_brut": mensuel_brut,
        "mensuel_net": mensuel_net,
    }


# ---------------------------------------------------------------------------
# Document
# ---------------------------------------------------------------------------

def generer_devis(devis, societe) -> BytesIO:
    """Genere un devis Word oriente client et retourne un buffer BytesIO."""
    doc = Document()
    setup_page(doc)
    force_arial(doc)

    s = _calcul_synthese(devis)

    _add_header(doc, devis, societe)
    spacer(doc, 4)
    _add_emetteur_meta(doc, devis, societe)
    spacer(doc, 4)
    _add_client_objet(doc, devis)
    spacer(doc, 8)

    _add_accroche(doc, devis)
    spacer(doc, 8)

    _add_creation(doc, devis, s)
    spacer(doc, 8)

    _add_socle(doc)
    spacer(doc, 8)

    if s["avantages_total"] > 0:
        _add_avantages(doc, devis, s)
        spacer(doc, 8)

    if s["mensuel_net"] > 0:
        _add_abonnement(doc, devis, s)
        spacer(doc, 8)

    _add_recap_financier(doc, devis, s)
    spacer(doc, 8)

    if _mode_label(devis) == "Leasing":
        _add_leasing(doc, devis)
    else:
        _add_echeancier(doc, devis)
    spacer(doc, 6)

    if devis.note_commerciale:
        _add_note(doc, devis.note_commerciale)
        spacer(doc, 4)

    _add_mentions(doc, devis, societe)

    buf = BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf


def _add_header(doc, devis, societe):
    """En-tete : logos FluXweb (gauche) + titre DEVIS (droite), filet navy.

    Favicon et wordmark sont places dans des cellules distinctes, centrees
    verticalement, pour qu'ils s'alignent sur leur centre (et non sur la
    ligne de base, qui ferait remonter le favicon plus haut).
    """
    mark = _ASSETS / "logo_mark.png"
    wordmark = _ASSETS / "logo_wordmark.png"
    has_logos = mark.exists() and wordmark.exists()

    tbl = doc.add_table(rows=1, cols=3)
    tbl_no_spacing(tbl)
    c_mark, c_word, c_titre = tbl.rows[0].cells
    cell_w(c_mark, 1.2)
    cell_w(c_word, 6.0)
    cell_w(c_titre, 10.8)
    for c in (c_mark, c_word, c_titre):
        c.vertical_alignment = WD_ALIGN_VERTICAL.CENTER

    if has_logos:
        p = c_mark.paragraphs[0]
        p_fmt(p, before=0, after=0)
        p.add_run().add_picture(str(mark), height=Cm(0.8))
        p = c_word.paragraphs[0]
        p_fmt(p, before=0, after=0)
        p.add_run().add_picture(str(wordmark), height=Cm(0.62))
    else:
        marque = (societe.marque if societe and societe.marque else None) or \
                 (societe.nom if societe else "FluXweb")
        p = c_word.paragraphs[0]
        p_fmt(p, before=0, after=0)
        run(p, marque, bold=True, size=16, color=C_NAVY)

    p = c_titre.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    p_fmt(p, before=0, after=0)
    run(p, "DEVIS", bold=True, size=22, color=C_NAVY)

    hline(doc)


def _add_emetteur_meta(doc, devis, societe):
    tbl = doc.add_table(rows=1, cols=2)
    tbl_no_spacing(tbl)
    left, right = tbl.rows[0].cells[0], tbl.rows[0].cells[1]
    cell_w(left, 9.5)
    cell_w(right, 8.5)

    lignes_emetteur = []
    if societe:
        lignes_emetteur = [
            (societe.nom, True),
            (societe.forme_juridique or "", False),
            (societe.adresse or "", False),
            (societe.cp_ville or "", False),
            (f"SIRET : {societe.siret}" if societe.siret else "", False),
            (f"TVA : {societe.tva_intracom}" if societe.tva_intracom else "", False),
        ]
    for text, bold in lignes_emetteur:
        if not text:
            continue
        p = left.add_paragraph()
        p_fmt(p, before=0, after=0)
        run(p, text, bold=bold, size=8, color=C_TEXT)

    meta = [
        ("Devis n°", devis.reference),
        ("Date d’émission", devis.date_emission.strftime("%d/%m/%Y")),
        ("Valable jusqu’au", devis.date_validite.strftime("%d/%m/%Y")),
        ("Mode de règlement", _mode_label(devis)),
    ]
    if devis.commercial:
        meta.append(("Commercial", devis.commercial))
    for label, val in meta:
        p = right.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        p_fmt(p, before=0, after=0)
        run(p, f"{label} : ", bold=True, size=8, color=C_TEXT)
        run(p, str(val), size=8, color=C_TEXT)


def _add_client_objet(doc, devis):
    # DESTINATAIRE (bloc adresse, aligne a gauche sous l'emetteur)
    p = doc.add_paragraph()
    p_fmt(p, before=0, after=1)
    run(p, "DESTINATAIRE", bold=True, size=7, color=C_NAVY)
    cp_ville = " ".join(x for x in [devis.client_cp, devis.client_ville] if x)
    for line in [
        devis.client_raison_sociale,
        devis.client_interlocuteur,
        devis.client_adresse,
        cp_ville,
        f"SIRET : {devis.client_siret}" if devis.client_siret else "",
        f"Tél : {devis.client_telephone}" if devis.client_telephone else "",
    ]:
        if line:
            p = doc.add_paragraph()
            p_fmt(p, before=0, after=0)
            run(p, line, size=8, color=C_TEXT)

    # OBJET : ligne pleine largeur, alignee a gauche (convention courrier commercial)
    spacer(doc, 4)
    p = doc.add_paragraph()
    p_fmt(p, before=0, after=0)
    run(p, "Objet : ", bold=True, size=9, color=C_NAVY)
    sous = devis.offre_nom + (f" ({devis.offre_type_site})" if devis.offre_type_site else "")
    run(p, "Création de votre site internet — " + sous, size=9, color=C_TEXT)


def _add_accroche(doc, devis):
    texte = (devis.accroche or "").strip() or (
        "Voici notre proposition pour la création de votre site internet, "
        "pensé pour renforcer votre visibilité et convertir vos visiteurs en clients."
    )
    p = doc.add_paragraph()
    p_fmt(p, before=0, after=0)
    run(p, texte, italic=True, size=9, color=C_TEXT)


def _items_offerts_designations(devis) -> set[str]:
    return {(a.designation or "").strip() for a in (devis.articles_offerts or [])}


def _add_socle(doc):
    """Socle commun a toutes les offres (reassurance : ce qui est toujours inclus)."""
    tbl_head = doc.add_table(rows=1, cols=1)
    tbl_no_spacing(tbl_head)
    cell_bg(tbl_head.rows[0].cells[0], HEX_NAVY)
    cell_text(tbl_head.rows[0].cells[0], "INCLUS DANS TOUTES NOS OFFRES",
              bold=True, size=10, color=C_WHITE)

    tbl = doc.add_table(rows=len(SOCLE_COMMUN), cols=2)
    tbl_no_spacing(tbl)
    full_tbl_borders(tbl)
    for i, (item, detail) in enumerate(SOCLE_COMMUN):
        cell_w(tbl.rows[i].cells[0], 5)
        cell_w(tbl.rows[i].cells[1], 13)
        cell_text(tbl.rows[i].cells[0], item, bold=True, size=8, color=C_NAVY)
        cell_text(tbl.rows[i].cells[1], detail, size=8, color=C_TEXT)


def _add_creation(doc, devis, s):
    """Section CREATION : titre offre + perimetre (sans prix) + prix unique."""
    offerts_noms = _items_offerts_designations(devis)

    # En-tete de section (bandeau navy)
    tbl_head = doc.add_table(rows=1, cols=1)
    tbl_no_spacing(tbl_head)
    cell_bg(tbl_head.rows[0].cells[0], HEX_NAVY)
    cell_text(tbl_head.rows[0].cells[0], "CRÉATION DE VOTRE SITE",
              bold=True, size=10, color=C_WHITE)

    # Corps : sous-titre + liste a puces + prix
    tbl = doc.add_table(rows=1, cols=1)
    tbl_no_spacing(tbl)
    full_tbl_borders(tbl)
    cell = tbl.rows[0].cells[0]
    cell_w(cell, 18)

    # Sous-titre (l'offre)
    p = cell.paragraphs[0]
    p_fmt(p, before=2, after=2)
    titre = f"Site {devis.offre_nom}"
    if devis.offre_type_site:
        titre += f" ({devis.offre_type_site})"
    run(p, titre, bold=True, size=9, color=C_NAVY)

    # Perimetre
    p = cell.add_paragraph()
    p_fmt(p, before=2, after=1)
    run(p, "Votre projet comprend :", bold=True, size=8, color=C_TEXT)

    # Construire la liste (sans prix) : prestations + options setup (payantes + incluses)
    puces = []
    for lg in sorted(devis.lignes or [], key=lambda x: x.ordre):
        nom = (lg.designation or "").strip()
        if nom:
            qte = lg.quantite or 1
            libelle = nom if qte <= 1 else f"{nom} (x{qte})"
            puces.append((libelle, "offert" if nom in offerts_noms else "normal"))
    for opt in sorted(devis.options or [], key=lambda x: x.ordre):
        type_ligne = (opt.type_ligne or "").upper()
        if type_ligne in ("RECURRENT", "OPTION_RECURRENT", "PACK"):
            continue  # va dans l'abonnement
        nom = (opt.nom or "").strip()
        if not nom:
            continue
        qte = opt.quantite or 1
        libelle = nom if qte <= 1 else f"{nom} (x{qte})"
        if nom in offerts_noms:
            etat = "offert"
        elif opt.inclus or _q(opt.prix_setup_ht) == 0:
            etat = "inclus"
        else:
            etat = "normal"
        puces.append((libelle, etat))

    for libelle, etat in puces:
        p = cell.add_paragraph()
        p_fmt(p, before=0, after=0)
        run(p, "•  ", size=8, color=C_NAVY)
        run(p, libelle, size=8, color=C_TEXT)
        if etat == "offert":
            run(p, "  (offert)", size=8, color=C_PAID, bold=True)
        elif etat == "inclus":
            run(p, "  (inclus)", size=8, color=C_AHEAD, italic=True)

    # Prix de la creation
    p = cell.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    p_fmt(p, before=4, after=2)
    run(p, "Prix de la création HT : ", bold=True, size=10, color=C_TEXT)
    run(p, fmt_eur(s["prix_creation_ht"]), bold=True, size=11, color=C_NAVY)


def _add_avantages(doc, devis, s):
    """Encadre VOS AVANTAGES : offerts (valeur) + remise + total."""
    rows = []
    for o in s["offerts"]:
        rows.append((f"Offert : {o['designation']}", "valeur " + fmt_eur(o["valeur"])))
    if s["remise_setup"] > 0:
        label_remise = "Remise commerciale"
        if s["remise_pct_setup"] > 0:
            label_remise += f" (− {_fmt_pct(s['remise_pct_setup'])})"
        rows.append((label_remise, "− " + fmt_eur(s["remise_setup"])))

    # En-tete
    tbl_head = doc.add_table(rows=1, cols=1)
    tbl_no_spacing(tbl_head)
    cell_bg(tbl_head.rows[0].cells[0], HEX_PAID)
    titre_av = f"VOS AVANTAGES — valables jusqu’au {devis.date_validite.strftime('%d/%m/%Y')}"
    cell_text(tbl_head.rows[0].cells[0], titre_av, bold=True, size=10, color=C_PAID)

    tbl = doc.add_table(rows=len(rows) + 1, cols=2)
    tbl_no_spacing(tbl)
    full_tbl_borders(tbl)
    for i, (label, val) in enumerate(rows):
        cell_bg(tbl.rows[i].cells[0], HEX_AVANTAGE)
        cell_bg(tbl.rows[i].cells[1], HEX_AVANTAGE)
        cell_w(tbl.rows[i].cells[0], 14)
        cell_w(tbl.rows[i].cells[1], 4)
        cell_text(tbl.rows[i].cells[0], label, size=8, color=C_PAID)
        cell_text(tbl.rows[i].cells[1], val, size=8, color=C_PAID,
                  align=WD_ALIGN_PARAGRAPH.RIGHT)

    # Total avantages (ligne mise en avant)
    last = len(rows)
    cell_bg(tbl.rows[last].cells[0], HEX_PAID)
    cell_bg(tbl.rows[last].cells[1], HEX_PAID)
    cell_w(tbl.rows[last].cells[0], 14)
    cell_w(tbl.rows[last].cells[1], 4)
    label_total = "Soit un total d’avantages offerts sur votre projet"
    if s["avantages_pct"] > 0:
        label_total += f" ({_fmt_pct(s['avantages_pct'])} de la valeur catalogue)"
    cell_text(tbl.rows[last].cells[0], label_total,
              bold=True, size=8, color=C_PAID)
    cell_text(tbl.rows[last].cells[1], fmt_eur(s["avantages_total"]),
              bold=True, size=9, color=C_PAID, align=WD_ALIGN_PARAGRAPH.RIGHT)


def _add_abonnement(doc, devis, s):
    """Section ABONNEMENT MENSUEL : perimetre recurrent + valorisation commerciale.

    Le detail comptable HT/TVA/TTC figure dans le RECAPITULATIF FINANCIER.
    """
    # En-tete
    tbl_head = doc.add_table(rows=1, cols=1)
    tbl_no_spacing(tbl_head)
    cell_bg(tbl_head.rows[0].cells[0], HEX_NAVY)
    cell_text(tbl_head.rows[0].cells[0], "ABONNEMENT MENSUEL",
              bold=True, size=10, color=C_WHITE)

    tbl = doc.add_table(rows=1, cols=1)
    tbl_no_spacing(tbl)
    full_tbl_borders(tbl)
    cell = tbl.rows[0].cells[0]
    cell_w(cell, 18)

    p = cell.paragraphs[0]
    p_fmt(p, before=2, after=1)
    run(p, "Comprend :", bold=True, size=8, color=C_TEXT)

    for opt in sorted(devis.options or [], key=lambda x: x.ordre):
        type_ligne = (opt.type_ligne or "").upper()
        if type_ligne not in ("RECURRENT", "OPTION_RECURRENT", "PACK"):
            continue
        nom = (opt.nom or "").strip()
        if not nom:
            continue
        p = cell.add_paragraph()
        p_fmt(p, before=0, after=0)
        run(p, "•  ", size=8, color=C_NAVY)
        run(p, nom, size=8, color=C_TEXT)
        if opt.inclus:
            run(p, "  (inclus)", size=8, color=C_AHEAD, italic=True)

    mensuel_net = s["mensuel_net"]
    mensuel_brut = s["mensuel_brut"]

    # Tarif catalogue barre (uniquement si avantage sur le recurrent)
    if mensuel_brut > mensuel_net:
        p = cell.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        p_fmt(p, before=4, after=0)
        run(p, "Tarif catalogue : ", size=8, color=C_AHEAD)
        run(p, fmt_eur(mensuel_brut) + " HT /mois", size=8, color=C_AHEAD, strike=True)

    # Votre tarif HT (avec % de reduction reel)
    p = cell.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    p_fmt(p, before=2 if mensuel_brut > mensuel_net else 4, after=0)
    run(p, "Votre tarif HT : ", bold=True, size=8, color=C_TEXT)
    run(p, fmt_eur(mensuel_net) + " /mois", bold=True, size=10, color=C_NAVY)
    if mensuel_brut > mensuel_net:
        reduction_pct = (mensuel_brut - mensuel_net) / mensuel_brut * D("100")
        run(p, f"  (− {_fmt_pct(reduction_pct.quantize(D('1'), ROUND_HALF_UP))})",
            bold=True, size=8, color=C_PAID)

    # Economie annualisee (levier fort : un petit mensuel devient un gros chiffre annuel)
    if mensuel_brut > mensuel_net:
        economie_an = _q((mensuel_brut - mensuel_net) * D("12"))
        p = cell.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        p_fmt(p, before=0, after=0)
        run(p, "soit " + fmt_eur(economie_an) + " d’économie par an", bold=True,
            size=8, color=C_PAID)

    # Note (le detail HT/TVA/TTC est dans le recapitulatif financier ci-dessous)
    p = doc.add_paragraph()
    p_fmt(p, before=4, after=0)
    run(p, "Démarre à la mise en ligne du site. Engagement initial de 12 mois, "
           "reconductible par tacite reconduction.",
        italic=True, size=7, color=C_AHEAD)


def _recap_sous_titre(doc, texte):
    """Sous-titre d'un bloc du recapitulatif financier."""
    p = doc.add_paragraph()
    p_fmt(p, before=4, after=2)
    run(p, texte, bold=True, size=8, color=C_NAVY)


def _recap_table(doc, lignes, suffix_total=""):
    """Petit tableau comptable a 2 colonnes (libelle / montant), derniere ligne = total navy."""
    tbl = doc.add_table(rows=len(lignes), cols=2)
    tbl_no_spacing(tbl)
    full_tbl_borders(tbl)
    for i, (label, val, is_total) in enumerate(lignes):
        cell_w(tbl.rows[i].cells[0], 14)
        cell_w(tbl.rows[i].cells[1], 4)
        if is_total:
            cell_bg(tbl.rows[i].cells[0], HEX_NAVY)
            cell_bg(tbl.rows[i].cells[1], HEX_NAVY)
        color = C_WHITE if is_total else C_TEXT
        size = 10 if is_total else 9
        cell_text(tbl.rows[i].cells[0], label, bold=is_total, size=size, color=color,
                  align=WD_ALIGN_PARAGRAPH.RIGHT)
        cell_text(tbl.rows[i].cells[1], val + (suffix_total if is_total else ""),
                  bold=is_total, size=size, color=color, align=WD_ALIGN_PARAGRAPH.RIGHT)


def _add_recap_financier(doc, devis, s):
    """Tableau RECAPITULATIF FINANCIER : creation (one-shot) puis abonnement mensuel.

    Vue purement comptable (HT -> TVA -> TTC). La valorisation des avantages reste
    dans l'encadre VOS AVANTAGES ; ce bloc ne fait que l'addition.
    """
    # En-tete de section (bandeau navy)
    tbl_head = doc.add_table(rows=1, cols=1)
    tbl_no_spacing(tbl_head)
    cell_bg(tbl_head.rows[0].cells[0], HEX_NAVY)
    cell_text(tbl_head.rows[0].cells[0], "RÉCAPITULATIF FINANCIER",
              bold=True, size=10, color=C_WHITE)

    # --- Bloc creation (one-shot) ---
    # total_ht est NET de remise ; le brut catalogue = net + remise.
    _recap_sous_titre(doc, "Création et mise en place (à régler à la commande)")
    net_ht = _q(devis.total_ht)
    if s["remise_setup"] > 0:
        brut_ht = net_ht + s["remise_setup"]
        label_remise = "Remise commerciale"
        if s["remise_pct_setup"] > 0:
            label_remise += f" (− {_fmt_pct(s['remise_pct_setup'])})"
        lignes = [
            ("Création et mise en place HT", fmt_eur(brut_ht), False),
            (label_remise, "− " + fmt_eur(s["remise_setup"]), False),
            ("Sous-total HT", fmt_eur(net_ht), False),
        ]
    else:
        lignes = [("Création et mise en place HT", fmt_eur(net_ht), False)]
    lignes.append(("TVA 20 %", fmt_eur(_q(devis.total_tva)), False))
    lignes.append(("TOTAL TTC À LA COMMANDE", fmt_eur(_q(devis.total_ttc)), True))
    _recap_table(doc, lignes, suffix_total=" TTC")

    # --- Bloc abonnement mensuel ---
    if s["mensuel_net"] > 0:
        spacer(doc, 4)
        _recap_sous_titre(doc, "Abonnement mensuel (maintenance & hébergement)")
        mensuel_net = s["mensuel_net"]
        mensuel_tva = _q(mensuel_net * TVA_RATE)
        mensuel_ttc = _q(mensuel_net * (D("1") + TVA_RATE))
        lignes = [
            ("Abonnement HT", fmt_eur(mensuel_net) + " /mois", False),
            ("TVA 20 %", fmt_eur(mensuel_tva) + " /mois", False),
            ("TOTAL TTC MENSUEL", fmt_eur(mensuel_ttc) + " /mois", True),
        ]
        _recap_table(doc, lignes)


def _add_echeancier(doc, devis):
    parts = repartition_echeances(_plan_label(devis), devis.total_ttc)
    if len(parts) <= 1:
        return  # 100 % : pas d'echeancier a detailler

    p = doc.add_paragraph()
    run(p, "Échéancier de paiement", bold=True, size=9, color=C_NAVY)

    tbl = doc.add_table(rows=1 + len(parts), cols=2)
    tbl_no_spacing(tbl)
    full_tbl_borders(tbl)
    for i, h in enumerate(["Échéance", "Montant TTC"]):
        cell_w(tbl.rows[0].cells[i], 14 if i == 0 else 4)
        cell_bg(tbl.rows[0].cells[i], HEX_NAVY)
        cell_text(tbl.rows[0].cells[i], h, bold=True, size=8, color=C_WHITE,
                  align=WD_ALIGN_PARAGRAPH.RIGHT if i == 1 else WD_ALIGN_PARAGRAPH.LEFT)

    for idx, (label, montant) in enumerate(parts, start=1):
        cell_w(tbl.rows[idx].cells[0], 14)
        cell_w(tbl.rows[idx].cells[1], 4)
        cell_text(tbl.rows[idx].cells[0], label, size=8)
        cell_text(tbl.rows[idx].cells[1], fmt_eur(montant), size=8,
                  align=WD_ALIGN_PARAGRAPH.RIGHT)


def _add_leasing(doc, devis):
    p = doc.add_paragraph()
    run(p, "Financement (location évolutive)", bold=True, size=9, color=C_NAVY)

    lignes = []
    if devis.loyer_mensuel:
        lignes.append(("Loyer mensuel HT", fmt_eur(_q(devis.loyer_mensuel)) + " /mois"))
    if devis.duree_financement_mois:
        lignes.append(("Durée du financement", f"{devis.duree_financement_mois} mois"))
    if devis.montant_finance:
        lignes.append(("Montant financé", fmt_eur(_q(devis.montant_finance))))
    if not lignes:
        return

    tbl = doc.add_table(rows=len(lignes), cols=2)
    tbl_no_spacing(tbl)
    full_tbl_borders(tbl)
    for i, (label, val) in enumerate(lignes):
        cell_w(tbl.rows[i].cells[0], 14)
        cell_w(tbl.rows[i].cells[1], 4)
        cell_text(tbl.rows[i].cells[0], label, bold=True, size=8, color=C_TEXT,
                  align=WD_ALIGN_PARAGRAPH.RIGHT)
        cell_text(tbl.rows[i].cells[1], val, size=8, color=C_TEXT,
                  align=WD_ALIGN_PARAGRAPH.RIGHT)


def _add_note(doc, note):
    p = doc.add_paragraph()
    p_fmt(p, before=0, after=1)
    run(p, "Note", bold=True, size=7, color=C_NAVY)
    p = doc.add_paragraph()
    p_fmt(p, before=0, after=0)
    run(p, note, size=8, color=C_TEXT)


def _add_mentions(doc, devis, societe):
    hline(doc)
    mentions = [
        f"Devis valable jusqu’au {devis.date_validite.strftime('%d/%m/%Y')}.",
        "Pour acceptation, retourner ce devis daté et signé avec la mention "
        "« Bon pour accord ».",
    ]
    if societe and societe.iban:
        mentions.append(
            f"Règlement par virement bancaire : IBAN {societe.iban}"
            + (f" — BIC {societe.bic}" if societe.bic else "")
        )
    mentions.append("Prix exprimés en euros. TVA au taux en vigueur (20 %).")
    for m in mentions:
        p = doc.add_paragraph()
        p_fmt(p, before=0, after=1)
        run(p, m, italic=True, size=7, color=C_AHEAD)

    # Zone de signature
    spacer(doc, 8)
    p = doc.add_paragraph()
    p_fmt(p, before=0, after=0)
    run(p, "Bon pour accord, le ......./......./...........", bold=True, size=8, color=C_TEXT)
    p = doc.add_paragraph()
    p_fmt(p, before=2, after=0)
    run(p, "Signature du client (précédée de la mention « Bon pour accord ») :",
        size=8, color=C_TEXT)


def _mode_label(devis) -> str:
    m = devis.mode_reglement
    return m.value if hasattr(m, "value") else str(m)


def _plan_label(devis) -> str:
    p = devis.plan_paiement
    if p is None:
        return "100%"
    return p.value if hasattr(p, "value") else str(p)
