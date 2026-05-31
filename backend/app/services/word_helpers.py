"""Helpers partages pour la generation de documents Word.

Extrait des fonctions dupliquees dans build_factures_acompte.py
et build_factures_maintenance.py du tarificateur.
"""

from decimal import Decimal
from pathlib import Path
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml

# Palette navy (identique au tarificateur)
C_NAVY = RGBColor(0x1A, 0x35, 0x5E)
C_BAND = RGBColor(0x2E, 0x5F, 0x9E)
C_TEXT = RGBColor(0x1A, 0x1A, 0x1A)
C_WHITE = RGBColor(0xFF, 0xFF, 0xFF)
C_ROW_A = RGBColor(0xF5, 0xF7, 0xFA)
C_SUBTOT = RGBColor(0xE8, 0xEC, 0xF2)
C_PAID = RGBColor(0x21, 0x7A, 0x3C)
C_CURR = RGBColor(0xC0, 0x60, 0x00)
C_AHEAD = RGBColor(0x70, 0x70, 0x70)

HEX_NAVY = "1A355E"
HEX_CURR = "FFF8EC"
HEX_PAID = "EEF7F1"
HEX_BORDER = "BDC7D6"

FONT_NAME = "Arial"
NNBSP = "\u202f"
EURO = "\u20ac"


def dxa(cm: float) -> str:
    return str(int(cm * 567))


def fmt_eur(v) -> str:
    """Formate un montant en euros : 1 234,56 EUR."""
    val = float(v)
    parts = f"{val:,.2f}".replace(",", NNBSP).replace(".", ",")
    return f"{parts}{NNBSP}{EURO}"


def setup_page(doc: Document):
    """Configure la page A4 avec marges standard."""
    for section in doc.sections:
        section.page_width = Cm(21)
        section.page_height = Cm(29.7)
        section.top_margin = Cm(1.8)
        section.bottom_margin = Cm(1.8)
        section.left_margin = Cm(1.5)
        section.right_margin = Cm(1.5)


def force_arial(doc: Document):
    """Force la police Arial sur tous les styles."""
    for style in doc.styles:
        try:
            rpr = style.element.get_or_add_rPr()
            for tag in (qn("w:rFonts"),):
                old = rpr.find(tag)
                if old is not None:
                    rpr.remove(old)
            fonts = parse_xml(
                f'<w:rFonts {nsdecls("w")} '
                f'w:ascii="{FONT_NAME}" w:hAnsi="{FONT_NAME}" '
                f'w:cs="{FONT_NAME}" w:eastAsia="{FONT_NAME}"/>'
            )
            rpr.append(fonts)
        except Exception:
            pass


def tbl_no_spacing(tbl):
    """Supprime l'espacement entre cellules."""
    props = tbl._tbl.tblPr
    if props is None:
        props = parse_xml(f"<w:tblPr {nsdecls('w')}/>")
        tbl._tbl.insert(0, props)
    spacing = parse_xml(f'<w:tblCellSpacing {nsdecls("w")} w:w="0" w:type="dxa"/>')
    props.append(spacing)


def full_tbl_borders(tbl, color: str = HEX_BORDER, sz: int = 4):
    """Applique des bordures completes a une table."""
    borders_xml = f"""<w:tblBorders {nsdecls("w")}>
        <w:top w:val="single" w:sz="{sz}" w:space="0" w:color="{color}"/>
        <w:left w:val="single" w:sz="{sz}" w:space="0" w:color="{color}"/>
        <w:bottom w:val="single" w:sz="{sz}" w:space="0" w:color="{color}"/>
        <w:right w:val="single" w:sz="{sz}" w:space="0" w:color="{color}"/>
        <w:insideH w:val="single" w:sz="{sz}" w:space="0" w:color="{color}"/>
        <w:insideV w:val="single" w:sz="{sz}" w:space="0" w:color="{color}"/>
    </w:tblBorders>"""
    tbl._tbl.tblPr.append(parse_xml(borders_xml))


def cell_bg(cell, hex_color: str):
    """Definit la couleur de fond d'une cellule."""
    shading = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{hex_color}" w:val="clear"/>')
    cell._tc.get_or_add_tcPr().append(shading)


def cell_w(cell, cm: float):
    """Definit la largeur d'une cellule."""
    width = parse_xml(f'<w:tcW {nsdecls("w")} w:w="{dxa(cm)}" w:type="dxa"/>')
    cell._tc.get_or_add_tcPr().append(width)


def row_height(row, cm: float, exact: bool = True):
    """Definit la hauteur d'une ligne."""
    rule = "exact" if exact else "atLeast"
    rpr = parse_xml(f'<w:trPr {nsdecls("w")}><w:trHeight w:val="{dxa(cm)}" w:hRule="{rule}"/></w:trPr>')
    row._tr.insert(0, rpr)


def p_fmt(para, before: float = 0, after: float = 0, line: float | None = None):
    """Configure l'espacement d'un paragraphe (en points)."""
    fmt = para.paragraph_format
    fmt.space_before = Pt(before)
    fmt.space_after = Pt(after)
    if line is not None:
        fmt.line_spacing = Pt(line)


def run(para, text: str, bold: bool = False, size: float = 9,
        color: RGBColor | None = None, italic: bool = False, strike: bool = False):
    """Ajoute un run formate a un paragraphe."""
    r = para.add_run(text)
    r.bold = bold
    r.font.size = Pt(size)
    r.font.name = FONT_NAME
    if color:
        r.font.color.rgb = color
    if italic:
        r.italic = True
    if strike:
        r.font.strike = True
    return r


def cell_text(cell, text: str, bold: bool = False, size: float = 9,
              color: RGBColor | None = None, align=WD_ALIGN_PARAGRAPH.LEFT,
              italic: bool = False, strike: bool = False,
              v_align=WD_ALIGN_VERTICAL.CENTER, before: float = 1, after: float = 1):
    """Ecrit du texte formate dans une cellule de tableau."""
    cell.vertical_alignment = v_align
    para = cell.paragraphs[0]
    para.alignment = align
    p_fmt(para, before=before, after=after)
    run(para, text, bold=bold, size=size, color=color, italic=italic, strike=strike)


def hline(doc: Document, color: str = HEX_NAVY, thickness: int = 6):
    """Ajoute une ligne horizontale."""
    para = doc.add_paragraph()
    p_fmt(para, before=2, after=2)
    pPr = para._p.get_or_add_pPr()
    borders = parse_xml(
        f'<w:pBdr {nsdecls("w")}>'
        f'<w:bottom w:val="single" w:sz="{thickness}" w:space="1" w:color="{color}"/>'
        f"</w:pBdr>"
    )
    pPr.append(borders)


def spacer(doc: Document, pt: float = 6):
    """Ajoute un espace vertical."""
    para = doc.add_paragraph()
    p_fmt(para, before=0, after=0, line=pt)


# ---------------------------------------------------------------------------
# Blocs d'en-tete communs aux documents (devis ET factures) — source unique
# de presentation pour eviter toute divergence de style entre documents.
# ---------------------------------------------------------------------------

ASSETS_DIR = Path(__file__).resolve().parent.parent / "templates" / "assets"


def add_logo_header(doc: Document, titre: str, sous_titre: str | None = None,
                    titre_size: float = 22, marque_fallback: str = "FluXweb"):
    """En-tete : logos FluXweb (gauche) + titre (droite) + filet navy.

    Favicon et wordmark sont dans des cellules distinctes, centrees
    verticalement, pour qu'ils s'alignent sur leur centre (et non sur la
    ligne de base, qui ferait remonter le favicon).
    """
    mark = ASSETS_DIR / "logo_mark.png"
    wordmark = ASSETS_DIR / "logo_wordmark.png"
    has_logos = mark.exists() and wordmark.exists()

    tbl = doc.add_table(rows=1, cols=3)
    tbl_no_spacing(tbl)
    c_mark, c_word, c_titre = tbl.rows[0].cells
    cell_w(c_mark, 1.2)
    cell_w(c_word, 6.0)
    cell_w(c_titre, 10.8)
    for c in (c_mark, c_word, c_titre):
        c.vertical_alignment = WD_ALIGN_VERTICAL.CENTER

    if has_logos:
        p = c_mark.paragraphs[0]
        p_fmt(p, before=0, after=0)
        p.add_run().add_picture(str(mark), height=Cm(0.8))
        p = c_word.paragraphs[0]
        p_fmt(p, before=0, after=0)
        p.add_run().add_picture(str(wordmark), height=Cm(0.62))
    else:
        p = c_word.paragraphs[0]
        p_fmt(p, before=0, after=0)
        run(p, marque_fallback, bold=True, size=16, color=C_NAVY)

    p = c_titre.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    p_fmt(p, before=0, after=0)
    run(p, titre, bold=True, size=titre_size, color=C_NAVY)
    if sous_titre:
        p2 = c_titre.add_paragraph()
        p2.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        p_fmt(p2, before=0, after=0)
        run(p2, sous_titre, bold=True, size=10, color=C_AHEAD)

    hline(doc)


def add_emetteur_meta(doc: Document, emetteur_lines, meta_rows):
    """Ligne emetteur (gauche) + metadonnees du document (droite).

    emetteur_lines : liste de (texte, bold). meta_rows : liste de (label, valeur).
    """
    tbl = doc.add_table(rows=1, cols=2)
    tbl_no_spacing(tbl)
    left, right = tbl.rows[0].cells[0], tbl.rows[0].cells[1]
    cell_w(left, 9.5)
    cell_w(right, 8.5)

    for text, bold in emetteur_lines:
        if not text:
            continue
        p = left.add_paragraph()
        p_fmt(p, before=0, after=0)
        run(p, text, bold=bold, size=8, color=C_TEXT)

    for label, val in meta_rows:
        p = right.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        p_fmt(p, before=0, after=0)
        run(p, f"{label} : ", bold=True, size=8, color=C_TEXT)
        run(p, str(val), size=8, color=C_TEXT)


def add_destinataire(doc: Document, lignes):
    """Bloc DESTINATAIRE (adresse client), aligne a gauche."""
    p = doc.add_paragraph()
    p_fmt(p, before=0, after=1)
    run(p, "DESTINATAIRE", bold=True, size=7, color=C_NAVY)
    for line in lignes:
        if line:
            p = doc.add_paragraph()
            p_fmt(p, before=0, after=0)
            run(p, line, size=8, color=C_TEXT)


def add_objet(doc: Document, texte: str):
    """Ligne 'Objet : ...' pleine largeur, alignee a gauche (convention)."""
    p = doc.add_paragraph()
    p_fmt(p, before=0, after=0)
    run(p, "Objet : ", bold=True, size=9, color=C_NAVY)
    run(p, texte, size=9, color=C_TEXT)
