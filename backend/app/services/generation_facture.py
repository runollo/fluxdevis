"""Service de generation de factures Word (acompte et maintenance).

Porte depuis build_factures_acompte.py et build_factures_maintenance.py.
"""

from datetime import date
from decimal import Decimal, ROUND_HALF_UP
from io import BytesIO

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH

from app.services.word_helpers import (
    C_NAVY, C_TEXT, C_WHITE, C_PAID, C_CURR, C_AHEAD,
    HEX_NAVY, HEX_CURR, HEX_PAID,
    fmt_eur, setup_page, force_arial, tbl_no_spacing, full_tbl_borders,
    cell_bg, cell_w, p_fmt, run, cell_text, hline, spacer,
    add_logo_header, add_emetteur_meta, add_destinataire, add_objet,
)

D = Decimal
TVA_RATE = D("0.20")


def _q(val) -> Decimal:
    return Decimal(str(val)).quantize(Decimal("0.01"), rounding=ROUND_HALF_UP)


class FactureData:
    """Donnees necessaires pour generer une facture."""

    def __init__(self, **kwargs):
        self.numero: str = kwargs.get("numero", "")
        self.type_facture: str = kwargs.get("type_facture", "acompte")
        self.date_emission: date = kwargs.get("date_emission", date.today())
        self.date_echeance: date = kwargs.get("date_echeance", date.today())
        self.objet: str = kwargs.get("objet", "")

        # Emetteur
        self.emetteur_nom: str = kwargs.get("emetteur_nom", "")
        self.emetteur_forme: str = kwargs.get("emetteur_forme", "")
        self.emetteur_marque: str = kwargs.get("emetteur_marque", "")
        self.emetteur_adresse: str = kwargs.get("emetteur_adresse", "")
        self.emetteur_cp_ville: str = kwargs.get("emetteur_cp_ville", "")
        self.emetteur_siret: str = kwargs.get("emetteur_siret", "")
        self.emetteur_rcs: str = kwargs.get("emetteur_rcs", "")
        self.emetteur_tva_num: str = kwargs.get("emetteur_tva_num", "")
        self.emetteur_email: str = kwargs.get("emetteur_email", "")
        self.emetteur_web: str = kwargs.get("emetteur_web", "")
        self.emetteur_iban: str = kwargs.get("emetteur_iban", "")
        self.emetteur_bic: str = kwargs.get("emetteur_bic", "")

        # Client
        self.client_societe: str = kwargs.get("client_societe", "")
        self.client_contact: str = kwargs.get("client_contact", "")
        self.client_adresse: str = kwargs.get("client_adresse", "")
        self.client_cp_ville: str = kwargs.get("client_cp_ville", "")
        self.client_siret: str = kwargs.get("client_siret", "")
        self.client_email: str = kwargs.get("client_email", "")

        # Lignes
        self.designation: str = kwargs.get("designation", "")
        self.quantite: str = kwargs.get("quantite", "1")
        self.prix_unitaire_ht: Decimal = D(str(kwargs.get("prix_unitaire_ht", 0)))
        self.montant_ht: Decimal = D(str(kwargs.get("montant_ht", 0)))

        # References
        self.devis_ref: str = kwargs.get("devis_ref", "")
        self.prestation: str = kwargs.get("prestation", "")
        self.periode: str | None = kwargs.get("periode")

        # Echeancier (acompte)
        self.echeances: list[dict] = kwargs.get("echeances", [])
        self.idx_echeance: int = kwargs.get("idx_echeance", 0)

    @property
    def montant_tva(self) -> Decimal:
        return _q(self.montant_ht * TVA_RATE)

    @property
    def montant_ttc(self) -> Decimal:
        return _q(self.montant_ht + self.montant_tva)


def generer_facture(data: FactureData) -> BytesIO:
    """Genere une facture Word et retourne un buffer BytesIO."""
    doc = Document()
    setup_page(doc)
    force_arial(doc)

    sous_map = {"maintenance": "de maintenance", "solde": "de solde", "acompte": "d\u2019acompte"}
    sous_titre = sous_map.get(data.type_facture, "d\u2019acompte")

    _add_header(doc, data, sous_titre)
    spacer(doc, 4)
    _add_emetteur_meta(doc, data)
    spacer(doc, 4)
    _add_client_objet(doc, data)
    spacer(doc, 6)
    _add_detail(doc, data)
    spacer(doc, 4)
    _add_totaux(doc, data)
    spacer(doc, 6)

    if data.type_facture == "acompte" and data.echeances:
        _add_echeancier(doc, data)
        spacer(doc, 6)

    _add_mentions(doc, data)

    buf = BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf


def _add_header(doc, data, sous_titre):
    marque = data.emetteur_marque or data.emetteur_nom or "FluXweb"
    add_logo_header(doc, "FACTURE", sous_titre=sous_titre, marque_fallback=marque)


def _add_emetteur_meta(doc, data):
    emetteur_lines = [
        (data.emetteur_nom, True),
        (data.emetteur_forme, False),
        (data.emetteur_adresse, False),
        (data.emetteur_cp_ville, False),
        (f"SIRET : {data.emetteur_siret}" if data.emetteur_siret else "", False),
        (f"RCS : {data.emetteur_rcs}" if data.emetteur_rcs else "", False),
        (f"TVA : {data.emetteur_tva_num}" if data.emetteur_tva_num else "", False),
    ]
    meta = [
        ("Facture n\u00b0", data.numero),
        ("Date d\u2019\u00e9mission", data.date_emission.strftime("%d/%m/%Y")),
        ("Date d\u2019\u00e9ch\u00e9ance", data.date_echeance.strftime("%d/%m/%Y")),
    ]
    if data.periode:
        meta.append(("P\u00e9riode", data.periode))
    add_emetteur_meta(doc, emetteur_lines, meta)


def _add_client_objet(doc, data):
    cp_ville = data.client_cp_ville
    add_destinataire(doc, [
        data.client_societe,
        data.client_contact,
        data.client_adresse,
        cp_ville,
        f"SIRET : {data.client_siret}" if data.client_siret else "",
    ])
    spacer(doc, 4)
    add_objet(doc, data.objet)


def _add_detail(doc, data):
    tbl = doc.add_table(rows=2, cols=5)
    tbl_no_spacing(tbl)
    full_tbl_borders(tbl)

    headers = ["D\u00e9signation", "Qt\u00e9", "P.U. HT", "TVA", "Montant HT"]
    widths = [8.0, 1.8, 3.0, 2.2, 3.0]
    for i, (h, w) in enumerate(zip(headers, widths)):
        cell_bg(tbl.rows[0].cells[i], HEX_NAVY)
        cell_w(tbl.rows[0].cells[i], w)
        cell_text(tbl.rows[0].cells[i], h, bold=True, size=8, color=C_WHITE,
                  align=WD_ALIGN_PARAGRAPH.CENTER)

    row = tbl.rows[1]
    vals = [data.designation, data.quantite, fmt_eur(data.prix_unitaire_ht),
            "20 %", fmt_eur(data.montant_ht)]
    aligns = [WD_ALIGN_PARAGRAPH.LEFT, WD_ALIGN_PARAGRAPH.CENTER,
              WD_ALIGN_PARAGRAPH.RIGHT, WD_ALIGN_PARAGRAPH.CENTER,
              WD_ALIGN_PARAGRAPH.RIGHT]
    for i, (v, a, w) in enumerate(zip(vals, aligns, widths)):
        cell_w(row.cells[i], w)
        cell_text(row.cells[i], v, size=8, align=a)


def _add_totaux(doc, data):
    tbl = doc.add_table(rows=4, cols=2)
    tbl_no_spacing(tbl)
    full_tbl_borders(tbl)

    lines = [
        ("Total HT", fmt_eur(data.montant_ht)),
        ("TVA 20 %", fmt_eur(data.montant_tva)),
        ("Total TTC", fmt_eur(data.montant_ttc)),
        ("Net \u00e0 payer", fmt_eur(data.montant_ttc)),
    ]
    for i, (label, val) in enumerate(lines):
        cell_w(tbl.rows[i].cells[0], 14)
        cell_w(tbl.rows[i].cells[1], 4)
        is_last = i == 3
        if is_last:
            cell_bg(tbl.rows[i].cells[0], HEX_NAVY)
            cell_bg(tbl.rows[i].cells[1], HEX_NAVY)
        color = C_WHITE if is_last else C_TEXT
        cell_text(tbl.rows[i].cells[0], label, bold=True, size=9, color=color,
                  align=WD_ALIGN_PARAGRAPH.RIGHT)
        cell_text(tbl.rows[i].cells[1], val, bold=is_last, size=9, color=color,
                  align=WD_ALIGN_PARAGRAPH.RIGHT)


def _add_echeancier(doc, data):
    p = doc.add_paragraph()
    run(p, "\u00c9ch\u00e9ancier de paiement", bold=True, size=9, color=C_NAVY)

    tbl = doc.add_table(rows=1 + len(data.echeances), cols=3)
    tbl_no_spacing(tbl)
    full_tbl_borders(tbl)

    ech_widths = [10.0, 4.0, 4.0]
    ech_aligns = [WD_ALIGN_PARAGRAPH.LEFT, WD_ALIGN_PARAGRAPH.CENTER,
                  WD_ALIGN_PARAGRAPH.RIGHT]
    for i, h in enumerate(["\u00c9ch\u00e9ance", "Date", "Montant TTC"]):
        cell_w(tbl.rows[0].cells[i], ech_widths[i])
        cell_bg(tbl.rows[0].cells[i], HEX_NAVY)
        cell_text(tbl.rows[0].cells[i], h, bold=True, size=8, color=C_WHITE,
                  align=ech_aligns[i])

    for row_idx, ech in enumerate(data.echeances, start=1):
        is_current = row_idx - 1 == data.idx_echeance
        is_paid = row_idx - 1 < data.idx_echeance
        row = tbl.rows[row_idx]
        if is_current:
            for c in row.cells:
                cell_bg(c, HEX_CURR)
        elif is_paid:
            for c in row.cells:
                cell_bg(c, HEX_PAID)
        for i in range(3):
            cell_w(row.cells[i], ech_widths[i])
        color = C_PAID if is_paid else (C_CURR if is_current else C_AHEAD)
        cell_text(row.cells[0], ech["label"], size=8, color=color, strike=is_paid)
        cell_text(row.cells[1], ech["date"], size=8, color=color, strike=is_paid,
                  align=WD_ALIGN_PARAGRAPH.CENTER)
        cell_text(row.cells[2], fmt_eur(ech["ttc"]), size=8, color=color,
                  strike=is_paid, align=WD_ALIGN_PARAGRAPH.RIGHT)


def _add_mentions(doc, data):
    hline(doc)
    mentions = [
        f"R\u00e8glement par virement bancaire : IBAN {data.emetteur_iban} \u2014 BIC {data.emetteur_bic}",
        "En cas de retard de paiement, une p\u00e9nalit\u00e9 de 3 fois le taux d\u2019int\u00e9r\u00eat l\u00e9gal sera appliqu\u00e9e.",
        "Indemnit\u00e9 forfaitaire pour frais de recouvrement : 40,00 \u20ac.",
        "Pas d\u2019escompte pour paiement anticip\u00e9.",
    ]
    if data.type_facture == "maintenance":
        mentions.append("Abonnement mensuel reconductible tacitement.")
    for m in mentions:
        p = doc.add_paragraph()
        p_fmt(p, before=0, after=1)
        run(p, m, italic=True, size=7, color=C_AHEAD)
